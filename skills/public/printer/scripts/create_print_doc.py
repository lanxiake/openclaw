#!/usr/bin/env python3
"""
Create Word document from images and text for printing
"""

import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

def create_print_document(output_path, items, title=None):
    """
    Create a Word document with images and text
    
    Args:
        output_path: Output .docx file path
        items: List of dicts with 'type' ('image' or 'text') and 'content' (path or text)
        title: Optional document title
    """
    doc = Document()
    
    # Add title if provided
    if title:
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    for item in items:
        if item['type'] == 'image':
            # Add image
            img_path = Path(item['content'])
            if not img_path.exists():
                print(f"Warning: Image not found: {img_path}", file=sys.stderr)
                continue
            
            # Get image dimensions to fit on page
            img = Image.open(img_path)
            width, height = img.size
            
            # Max width: 6 inches (leaving margins)
            max_width = 6.0
            if width > height:
                doc.add_picture(str(img_path), width=Inches(max_width))
            else:
                # Portrait: scale to fit
                aspect = height / width
                if aspect * max_width > 8:  # Too tall
                    doc.add_picture(str(img_path), height=Inches(8))
                else:
                    doc.add_picture(str(img_path), width=Inches(max_width))
            
            # Add caption if provided
            if 'caption' in item:
                caption = doc.add_paragraph(item['caption'])
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_format = caption.runs[0].font
                caption_format.size = Pt(10)
                caption_format.italic = True
            
            doc.add_paragraph()  # Spacing
            
        elif item['type'] == 'text':
            # Add text
            text = item['content']
            p = doc.add_paragraph(text)
            
            # Apply formatting if provided
            if 'style' in item:
                if item['style'] == 'heading':
                    p.style = 'Heading 2'
                elif item['style'] == 'bold':
                    p.runs[0].font.bold = True
    
    # Save document
    doc.save(output_path)
    print(f"✓ Created document: {output_path}")
    return 0

def main():
    parser = argparse.ArgumentParser(description='Create printable Word document')
    parser.add_argument('--output', '-o', required=True, help='Output .docx file')
    parser.add_argument('--title', '-t', help='Document title')
    parser.add_argument('--image', '-i', action='append', help='Add image (can be used multiple times)')
    parser.add_argument('--text', '-x', action='append', help='Add text (can be used multiple times)')
    
    args = parser.parse_args()
    
    items = []
    
    # Add images
    if args.image:
        for img_path in args.image:
            items.append({'type': 'image', 'content': img_path})
    
    # Add text
    if args.text:
        for text in args.text:
            items.append({'type': 'text', 'content': text})
    
    if not items:
        print("Error: No content provided (use --image or --text)", file=sys.stderr)
        return 1
    
    return create_print_document(args.output, items, args.title)

if __name__ == '__main__':
    sys.exit(main())
